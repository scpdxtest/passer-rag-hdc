"""
Convert paper_PaSSER_RAG_BdKCSE.md to an IEEE-conference-style .docx.

Differences vs. the MDPI converter (md_to_docx.py in the same folder):
- Times New Roman 10 pt body (vs. Palatino).
- Centered title (24 pt) and centered author block (10 pt italic).
- Single-column for title + authors + abstract + keywords;
  two-column body starting at the marker `---COLUMN-BREAK---`
  (which is consumed, not rendered).
- Roman-numeral section headings rendered in SMALL CAPS, centered.
- Lettered subsection headings (A., B., ...) rendered italic, left-aligned.
- IEEE-style reference paragraphs: a hanging indent so the [N] sticks to
  the left margin and the wrapped text aligns under the title.

Run from anywhere — paths are resolved relative to this script's location
(this file lives at newRAG/tools/, the paper at newRAG/). Pass --src and
--dst to override.

Setup (first time, on macOS Homebrew Python):
    pip3 install --user --break-system-packages python-docx
"""
import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# newRAG/tools/md_to_docx_ieee.py  → newRAG/
PAPER_DIR = Path(__file__).resolve().parent.parent

_argp = argparse.ArgumentParser(description=__doc__)
_argp.add_argument("--src", type=Path,
                   default=PAPER_DIR / "paper_PaSSER_RAG_BdKCSE.md",
                   help="Markdown source")
_argp.add_argument("--dst", type=Path,
                   default=PAPER_DIR / "paper_PaSSER_RAG_BdKCSE.docx",
                   help="DOCX destination")
_args = _argp.parse_args()
SRC = _args.src
DST = _args.dst


# --------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------

INLINE_RE = re.compile(
    r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
)

# Headings 1–4 + body all in Times New Roman by IEEE convention.
BODY_FONT = "Times New Roman"
BODY_SIZE = 10


def add_runs(paragraph, text, base_font=BODY_FONT, base_size=BODY_SIZE):
    """Append text to `paragraph` honouring ***bold-italic***, **bold**,
    *italic*, `code` runs. Newer markers take priority over shorter ones."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.name = base_font
            r.font.size = Pt(base_size)
        token = m.group(0)
        if token.startswith('***'):
            r = paragraph.add_run(token[3:-3])
            r.bold = True
            r.italic = True
        elif token.startswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(base_size - 1)
            pos = m.end()
            continue
        r.font.name = base_font
        r.font.size = Pt(base_size)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.name = base_font
        r.font.size = Pt(base_size)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def insert_continuous_section_break(doc, n_cols):
    """Insert a CONTINUOUS section break and configure the new section
    to use n_cols columns. python-docx doesn't expose this directly, so
    we build the XML by hand."""
    # Add a paragraph carrying the section break properties.
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    sect_pr = OxmlElement('w:sectPr')
    # Continuous break
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'continuous')
    sect_pr.append(type_el)
    # Page size + margins inherited from previous section by default,
    # but we re-state the columns count for the section that ENDS here.
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')   # the BEFORE section is 1-column
    cols.set(qn('w:space'), '720')
    sect_pr.append(cols)
    p_pr.append(sect_pr)
    # The NEXT section (everything after this paragraph) will use n_cols.
    # We set that on the document's *final* sectPr below; see save().
    doc._ieee_after_break_cols = n_cols


def finalise_columns(doc):
    """Set the document's terminal sectPr columns count to the value
    captured by insert_continuous_section_break(). Called just before
    saving."""
    n_cols = getattr(doc, '_ieee_after_break_cols', 1)
    sect_pr = doc.sections[-1]._sectPr
    # Replace or insert <w:cols>.
    for existing in sect_pr.findall(qn('w:cols')):
        sect_pr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(n_cols))
    cols.set(qn('w:space'), '432')      # ~0.3" between columns
    if n_cols > 1:
        cols.set(qn('w:equalWidth'), '1')
    sect_pr.append(cols)


# --------------------------------------------------------------------
# Document setup
# --------------------------------------------------------------------

doc = Document()

# IEEE conference margins: top/bottom 0.75", left/right 0.625" (approx).
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.625)
    section.right_margin  = Inches(0.625)

body_style = doc.styles['Normal']
body_style.font.name = BODY_FONT
body_style.font.size = Pt(BODY_SIZE)
body_style.paragraph_format.space_after = Pt(3)


def style_heading(name, size_pt, *, bold=True, italic=False, all_caps=False,
                  align=None, color=None, space_before=8, space_after=4):
    style = doc.styles[name]
    style.font.name = BODY_FONT
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic
    if all_caps:
        # python-docx doesn't expose font.all_caps directly; set via XML.
        rPr = style.element.get_or_add_rPr()
        caps = OxmlElement('w:caps')
        caps.set(qn('w:val'), '1')
        rPr.append(caps)
    if color:
        style.font.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = True
    if align is not None:
        pf.alignment = align


style_heading('Heading 1', 10, all_caps=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
style_heading('Heading 2', 10, bold=False, italic=True,
              align=WD_ALIGN_PARAGRAPH.LEFT)
style_heading('Heading 3', 10, bold=False, italic=True,
              align=WD_ALIGN_PARAGRAPH.LEFT)


# --------------------------------------------------------------------
# Parser
# --------------------------------------------------------------------

raw = SRC.read_text(encoding='utf-8').splitlines()
i = 0
N = len(raw)


def is_separator(s):
    return bool(re.fullmatch(r'\s*\|?(\s*:?-+:?\s*\|)+\s*', s.strip()))


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


seen_title = False
column_break_done = False

while i < N:
    line = raw[i]
    stripped = line.rstrip()

    # ---- Column-break marker (consumed, not rendered) ----
    if stripped == '---COLUMN-BREAK---':
        insert_continuous_section_break(doc, 2)
        column_break_done = True
        i += 1
        continue

    # ---- Plain horizontal rule (renders as a thin border) ----
    if stripped == '---':
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        i += 1
        continue

    # ---- Title (first H1) ----
    if stripped.startswith('# '):
        title_text = stripped[2:].strip()
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(title_text)
        r.bold = False
        r.font.size = Pt(24)
        r.font.name = BODY_FONT
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after  = Pt(18)
        seen_title = True
        i += 1
        continue

    # ---- Section headings ----
    if stripped.startswith('## '):
        doc.add_heading(stripped[3:].strip(), level=1)
        i += 1
        continue

    if stripped.startswith('### '):
        doc.add_heading(stripped[4:].strip(), level=2)
        i += 1
        continue

    if stripped.startswith('#### '):
        doc.add_heading(stripped[5:].strip(), level=3)
        i += 1
        continue

    # ---- Table ----
    if line.strip().startswith('|') and i + 1 < N and is_separator(raw[i + 1]):
        header = parse_table_row(line)
        col_count = len(header)
        j = i + 2
        rows = []
        while j < N and raw[j].strip().startswith('|'):
            rows.append(parse_table_row(raw[j]))
            j += 1
        table = doc.add_table(rows=1 + len(rows), cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for k, cell_text in enumerate(header):
            cell = table.rows[0].cells[k]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, cell_text, base_size=8)
            for run in p.runs:
                run.bold = True
            set_cell_shading(cell, 'E1E4E8')
        for ridx, row in enumerate(rows, start=1):
            for k in range(col_count):
                cell = table.rows[ridx].cells[k]
                cell.text = ''
                p = cell.paragraphs[0]
                add_runs(p, row[k] if k < len(row) else '', base_size=8)
        try:
            table.style = 'Light Grid Accent 1'
        except KeyError:
            pass
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        i = j
        continue

    # ---- References paragraph (hanging indent) ----
    m = re.match(r'^\s*\[(\d+)\]\s+(.*)$', line)
    if m:
        # Accumulate continuation lines (Markdown allows soft-wrapped
        # references but here we treat each ref as one paragraph).
        text = m.group(0).lstrip()
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        pf.space_after = Pt(3)
        add_runs(p, text, base_size=9)
        i += 1
        continue

    # ---- Ordered list ----
    m = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
    if m:
        text = m.group(2)
        p = doc.add_paragraph(style='List Number')
        add_runs(p, text)
        i += 1
        continue

    # ---- Unordered list ----
    m = re.match(r'^\s*[-*]\s+(.*)$', line)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, text)
        i += 1
        continue

    # ---- Blank line ----
    if stripped == '':
        i += 1
        continue

    # ---- Paragraph (accumulate consecutive non-blank lines) ----
    para_lines = [stripped]
    j = i + 1
    while j < N:
        nxt = raw[j].rstrip()
        if (nxt == '' or nxt == '---' or nxt == '---COLUMN-BREAK---'
                or nxt.startswith('#') or nxt.startswith('|')
                or re.match(r'^\s*\[(\d+)\]\s+', nxt)
                or re.match(r'^\s*(\d+)\.\s+', nxt)
                or re.match(r'^\s*[-*]\s+', nxt)):
            break
        para_lines.append(nxt)
        j += 1
    text = ' '.join(para_lines)

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(3)

    # Before the column break, the author block is centered with each
    # affiliation italicised on its own line. After the break, body
    # paragraphs are justified.
    if not column_break_done:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Mostly italic in the author block, but **Name** is rendered
        # bold by add_runs already.
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0.2)
    add_runs(p, text)
    i = j


# --------------------------------------------------------------------
# Save
# --------------------------------------------------------------------

finalise_columns(doc)
doc.save(str(DST))
print(f"wrote {DST}")
