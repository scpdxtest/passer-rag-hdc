"""
Convert paper_PaSSER_RAG.md to .docx using python-docx.

The converter is tuned for the specific markdown features used in this
manuscript: ATX headings, bold/italic inline, pipe tables, fenced code
blocks, ordered/unordered lists, horizontal rules. It does NOT aim to
be a general-purpose markdown engine.

Run from anywhere — paths are resolved relative to the script's
location (this file lives at newRAG/tools/, the paper at newRAG/).
Pass --src / --dst to override.

Setup (first time, on macOS Homebrew Python):
    pip3 install --user --break-system-packages python-docx
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# newRAG/tools/md_to_docx.py  → newRAG/
PAPER_DIR = Path(__file__).resolve().parent.parent

_argp = argparse.ArgumentParser(description=__doc__)
_argp.add_argument("--src", type=Path, default=PAPER_DIR / "paper_PaSSER_RAG.md",
                   help="Markdown source (default: ../paper_PaSSER_RAG.md relative to this script)")
_argp.add_argument("--dst", type=Path, default=PAPER_DIR / "paper_PaSSER_RAG.docx",
                   help="DOCX destination (default: ../paper_PaSSER_RAG.docx relative to this script)")
_args = _argp.parse_args()
SRC = _args.src
DST = _args.dst


# --------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------

INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')


def add_runs(paragraph, text):
    """Append text to `paragraph` honouring **bold**, *italic*, `code` runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


# --------------------------------------------------------------------
# Document setup
# --------------------------------------------------------------------

doc = Document()

# Page margins ~ MDPI Electronics single column ish
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# Body style
body_style = doc.styles['Normal']
body_style.font.name = 'Palatino Linotype'
body_style.font.size = Pt(10)


def set_heading_style(name, size_pt, bold=True, color=None, space_before=12, space_after=6):
    style = doc.styles[name]
    style.font.name = 'Palatino Linotype'
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.keep_with_next = True


set_heading_style('Heading 1', 14, color=RGBColor(0x00, 0x00, 0x00))
set_heading_style('Heading 2', 12, color=RGBColor(0x00, 0x00, 0x00))
set_heading_style('Heading 3', 11, color=RGBColor(0x00, 0x00, 0x00))
set_heading_style('Heading 4', 10, color=RGBColor(0x33, 0x33, 0x33))


# --------------------------------------------------------------------
# Parser
# --------------------------------------------------------------------

raw = SRC.read_text(encoding='utf-8').splitlines()
i = 0
N = len(raw)


def is_separator(s):
    return bool(re.fullmatch(r'\s*\|?(\s*:?-+:?\s*\|)+\s*', s.strip()))


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


while i < N:
    line = raw[i]
    stripped = line.rstrip()

    # ---- Horizontal rule ----
    if stripped == '---':
        # Use a paragraph border as a HR.
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        i += 1
        continue

    # ---- Headings ----
    if stripped.startswith('# '):
        # Title — center, larger.
        title_text = stripped[2:].strip()
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run(title_text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.name = 'Palatino Linotype'
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after  = Pt(12)
        i += 1
        continue

    if stripped.startswith('## '):
        doc.add_heading(stripped[3:].strip(), level=1)
        i += 1
        continue

    if stripped.startswith('### '):
        doc.add_heading(stripped[4:].strip(), level=2)
        i += 1
        continue

    if stripped.startswith('#### '):
        doc.add_heading(stripped[5:].strip(), level=3)
        i += 1
        continue

    # ---- Code block ----
    if stripped.startswith('```'):
        i += 1
        code_lines = []
        while i < N and not raw[i].rstrip().startswith('```'):
            code_lines.append(raw[i])
            i += 1
        # Skip the closing ```
        if i < N:
            i += 1
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        # Light gray shading via paragraph border
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F6F8FA')
        p_pr.append(shd)
        continue

    # ---- Table ----
    if line.strip().startswith('|') and i + 1 < N and is_separator(raw[i + 1]):
        # Collect all consecutive table lines.
        header = parse_table_row(line)
        # Skip the separator
        sep = raw[i + 1]
        col_count = len(header)
        j = i + 2
        rows = []
        while j < N and raw[j].strip().startswith('|'):
            rows.append(parse_table_row(raw[j]))
            j += 1
        # Build the docx table.
        table = doc.add_table(rows=1 + len(rows), cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        for k, cell_text in enumerate(header):
            cell = table.rows[0].cells[k]
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
            set_cell_shading(cell, 'E1E4E8')
        # Body rows
        for ridx, row in enumerate(rows, start=1):
            for k in range(col_count):
                cell = table.rows[ridx].cells[k]
                cell.text = ''
                p = cell.paragraphs[0]
                add_runs(p, row[k] if k < len(row) else '')
                for run in p.runs:
                    run.font.size = Pt(9)
        # Apply table grid
        table.style = 'Light Grid Accent 1'
        # Spacing after table
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        i = j
        continue

    # ---- Ordered list ----
    m = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
    if m:
        text = m.group(2)
        p = doc.add_paragraph(style='List Number')
        add_runs(p, text)
        i += 1
        continue

    # ---- Unordered list ----
    m = re.match(r'^\s*[-*]\s+(.*)$', line)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, text)
        i += 1
        continue

    # ---- Blank line ----
    if stripped == '':
        i += 1
        continue

    # ---- Paragraph ----
    # Accumulate consecutive non-blank, non-special lines as one paragraph.
    para_lines = [stripped]
    j = i + 1
    while j < N:
        nxt = raw[j].rstrip()
        if (nxt == '' or nxt == '---' or nxt.startswith('#')
                or nxt.startswith('```') or nxt.startswith('|')
                or re.match(r'^\s*(\d+)\.\s+', nxt) or re.match(r'^\s*[-*]\s+', nxt)):
            break
        para_lines.append(nxt)
        j += 1
    text = ' '.join(para_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)
    i = j

# --------------------------------------------------------------------
# Save
# --------------------------------------------------------------------

doc.save(str(DST))
print(f"wrote {DST}")
